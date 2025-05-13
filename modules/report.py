import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import io
import base64
from datetime import datetime
import numpy as np
from database_helper import get_all_questions, get_user_submissions, get_all_users
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

def format_date(date_value):
    """Định dạng ngày tháng từ nhiều kiểu dữ liệu khác nhau"""
    if not date_value:
        return "N/A"
    
    try:
        # Nếu là số nguyên (timestamp)
        if isinstance(date_value, (int, float)):
            return datetime.fromtimestamp(date_value).strftime("%d/%m/%Y")
        
        # Nếu là chuỗi ISO (từ Supabase)
        elif isinstance(date_value, str):
            try:
                # Thử parse chuỗi ISO
                dt = datetime.fromisoformat(date_value.replace('Z', '+00:00'))
                return dt.strftime("%d/%m/%Y")
            except:
                # Nếu không phải ISO, trả về nguyên bản
                return date_value
        
        # Nếu đã là đối tượng datetime
        elif isinstance(date_value, datetime):
            return date_value.strftime("%d/%m/%Y")
            
        # Các trường hợp khác, trả về dạng chuỗi
        else:
            return str(date_value)
    except Exception as e:
        print(f"Error formatting date: {e}, value type: {type(date_value)}, value: {date_value}")
        return "N/A"

def dataframe_to_docx(df, title, filename):
    """Tạo file DOCX từ DataFrame"""
    doc = Document()
    
    # Thiết lập font chữ mặc định
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Thêm tiêu đề
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thêm thời gian xuất báo cáo
    time_paragraph = doc.add_paragraph(f"Thời gian xuất báo cáo: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    time_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Tạo bảng
    # Thêm một hàng cho tiêu đề cột
    table = doc.add_table(rows=1, cols=len(df.columns), style='Table Grid')
    
    # Thêm tiêu đề cột
    header_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        header_cells[i].text = str(col_name)
        # Đặt kiểu cho tiêu đề
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0]
            run.bold = True
    
    # Thêm dữ liệu
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
            # Căn giữa cho các ô
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thêm chân trang
    doc.add_paragraph()
    footer = doc.add_paragraph("Hệ thống Khảo sát & Đánh giá")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Lưu tệp
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def create_student_report_docx(student_name, student_email, student_class, submission, questions, max_possible):
    """Tạo báo cáo chi tiết bài làm của học viên dạng DOCX"""
    doc = Document()
    
    # Thiết lập font chữ mặc định
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Thêm tiêu đề
    heading = doc.add_heading(f"Báo cáo chi tiết bài làm - {student_name}", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thêm thông tin học viên
    doc.add_heading("Thông tin học viên", level=2)
    info_table = doc.add_table(rows=4, cols=2, style='Table Grid')
    
    # Thêm dữ liệu vào bảng thông tin
    cells = info_table.rows[0].cells
    cells[0].text = "Họ và tên"
    cells[1].text = student_name
    
    cells = info_table.rows[1].cells
    cells[0].text = "Email"
    cells[1].text = student_email
    
    cells = info_table.rows[2].cells
    cells[0].text = "Lớp"
    cells[1].text = student_class
    
    cells = info_table.rows[3].cells
    cells[0].text = "Thời gian nộp"
    cells[1].text = datetime.fromtimestamp(submission["timestamp"]).strftime("%H:%M:%S %d/%m/%Y")
    
    # Tính toán thông tin về bài làm
    total_correct = 0
    total_questions = len(questions)
    
    doc.add_heading("Chi tiết câu trả lời", level=2)
    
    # Tạo bảng chi tiết câu trả lời
    answers_table = doc.add_table(rows=1, cols=5, style='Table Grid')
    
    # Thêm tiêu đề cho bảng
    header_cells = answers_table.rows[0].cells
    headers = ["Câu hỏi", "Đáp án của học viên", "Đáp án đúng", "Kết quả", "Điểm"]
    
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
    
    # Thêm dữ liệu câu trả lời
    for q in questions:
        q_id = str(q["id"])
        
        # Đáp án người dùng
        user_ans = submission["responses"].get(q_id, [])
        expected = [q["answers"][i - 1] for i in q["correct"]]
        
        # Kiểm tra đúng/sai
        is_correct = set(user_ans) == set(expected)
        if is_correct:
            total_correct += 1
            result = "Đúng"
            points = q["score"]
        else:
            result = "Sai"
            points = 0
        
        # Thêm hàng mới vào bảng
        row_cells = answers_table.add_row().cells
        
        # Thêm thông tin câu hỏi
        row_cells[0].text = f"Câu {q['id']}: {q['question']}"
        row_cells[1].text = ", ".join(user_ans) if user_ans else "Không trả lời"
        row_cells[2].text = ", ".join(expected)
        row_cells[3].text = result
        
        # Đặt màu cho kết quả
        for paragraph in row_cells[3].paragraphs:
            run = paragraph.runs[0]
            if is_correct:
                run.font.color.rgb = RGBColor(0, 128, 0)  # Màu xanh lá cho đúng
                run.bold = True
            else:
                run.font.color.rgb = RGBColor(255, 0, 0)  # Màu đỏ cho sai
                run.bold = True
        
        row_cells[4].text = str(points)
    
    # Thêm tổng kết
    doc.add_heading("Tổng kết", level=2)
    summary_table = doc.add_table(rows=3, cols=2, style='Table Grid')
    
    cells = summary_table.rows[0].cells
    cells[0].text = "Số câu đúng"
    cells[1].text = f"{total_correct}/{total_questions}"
    
    cells = summary_table.rows[1].cells
    cells[0].text = "Điểm số"
    cells[1].text = f"{submission['score']}/{max_possible}"
    
    cells = summary_table.rows[2].cells
    cells[0].text = "Tỷ lệ đúng"
    cells[1].text = f"{(total_correct/total_questions*100):.1f}%"
    
    # Thêm chân trang
    doc.add_paragraph()
    footer = doc.add_paragraph("Xuất báo cáo từ Hệ thống Khảo sát & Đánh giá")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    time_footer = doc.add_paragraph(f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    time_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Lưu tệp
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def get_download_link_docx(buffer, filename, text):
    """Tạo link tải xuống cho file DOCX"""
    b64 = base64.b64encode(buffer.getvalue()).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">📥 {text}</a>'
    return href

def export_to_excel(dataframes, sheet_names, filename):
    """Tạo file Excel với nhiều sheet từ các DataFrame"""
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        for df, sheet_name in zip(dataframes, sheet_names):
            df.to_excel(writer, sheet_name=sheet_name, index=False)
    
    data = output.getvalue()
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="{filename}">📥 {filename}</a>'
    return href

def view_statistics():
    st.title("📊 Báo cáo & thống kê")
    
    # Lấy dữ liệu từ database
    questions = get_all_questions()
    submissions = get_user_submissions()
    students = get_all_users(role="Học viên")
    
    if not questions:
        st.warning("Chưa có dữ liệu câu hỏi nào trong hệ thống.")
        return
    
    if not submissions:
        st.warning("Chưa có ai nộp khảo sát.")
        return
    
    # Tạo tab thống kê
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["Tổng quan", "Theo học viên", "Theo câu hỏi", "Danh sách học viên", "Xuất báo cáo"])
    
    # --- Dữ liệu chung ---
    max_possible = sum([q["score"] for q in questions])
    
    # Chuẩn bị DataFrame cho báo cáo
    all_submission_data = []
    for s in submissions:
        # Tìm thông tin học viên
        student_info = next((student for student in students if student["email"] == s["user_email"]), None)
        full_name = student_info["full_name"] if student_info else "Không xác định"
        class_name = student_info["class"] if student_info else "Không xác định"
        
        # Thêm thông tin cơ bản
        submission_data = {
            "ID": s["id"],
            "Email": s["user_email"],
            "Họ và tên": full_name,
            "Lớp": class_name,
            "Thời gian nộp": datetime.fromtimestamp(s["timestamp"]).strftime("%d/%m/%Y %H:%M:%S"),
            "Điểm số": s["score"],
            "Điểm tối đa": max_possible,
            "Tỷ lệ đúng": f"{(s['score']/max_possible*100):.1f}%"
        }
        
        # Thêm câu trả lời của từng câu hỏi
        for q in questions:
            q_id = str(q["id"])
            user_ans = s["responses"].get(q_id, [])
            expected = [q["answers"][i - 1] for i in q["correct"]]
            is_correct = set(user_ans) == set(expected)
            
            # Thêm thông tin câu hỏi
            submission_data[f"Câu {q_id}: {q['question']}"] = ", ".join(user_ans) if user_ans else "Không trả lời"
            submission_data[f"Câu {q_id} - Đúng/Sai"] = "Đúng" if is_correct else "Sai"
        
        all_submission_data.append(submission_data)
    
    # DataFrame chứa tất cả bài nộp
    df_all_submissions = pd.DataFrame(all_submission_data)
    
    with tab1:
        st.subheader("Tổng quan kết quả")
        
        # Thống kê cơ bản
        total_submissions = len(submissions)
        avg_score = sum([s["score"] for s in submissions]) / total_submissions if total_submissions > 0 else 0
        max_score = max([s["score"] for s in submissions]) if submissions else 0
        total_users = len(set([s["user_email"] for s in submissions]))
        
        # Hiển thị metrics
        col1, col2, col3 = st.columns(3)
        col1.metric("📝 Tổng số bài nộp", total_submissions)
        col1.metric("👥 Số học viên đã làm", total_users)
        
        col2.metric("📊 Điểm trung bình", f"{avg_score:.2f}/{max_possible}")
        col2.metric("🏆 Điểm cao nhất", f"{max_score}/{max_possible}")
        
        col3.metric("📋 Số câu hỏi", len(questions))
        col3.metric("👨‍🎓 Tổng số học viên", len(students))
        
        # Biểu đồ điểm số theo thời gian
        st.subheader("Điểm số theo thời gian")
        
        # Chuẩn bị dữ liệu
        time_data = []
        for s in submissions:
            time_data.append({
                "timestamp": datetime.fromtimestamp(s["timestamp"]),
                "score": s["score"],
                "user": s["user_email"]
            })
        
        if time_data:
            df_time = pd.DataFrame(time_data)
            df_time = df_time.sort_values("timestamp")
            
            # Vẽ biểu đồ
            fig, ax = plt.subplots(figsize=(10, 5))
            ax.plot(df_time["timestamp"], df_time["score"], marker='o')
            ax.set_ylabel("Điểm số")
            ax.set_xlabel("Thời gian nộp bài")
            ax.grid(True, linestyle='--', alpha=0.7)
            
            # Giảm số lượng tick trên trục x
            max_ticks = 6
            if len(df_time) > max_ticks:
                stride = len(df_time) // max_ticks
                plt.xticks(df_time["timestamp"][::stride])
            
            plt.tight_layout()  # Đảm bảo không bị cắt chữ
            st.pyplot(fig)
        
        # Hiển thị phân phối điểm
        st.subheader("Phân phối điểm số")
        if submissions:
            scores = [s["score"] for s in submissions]
            fig, ax = plt.subplots(figsize=(10, 5))
            ax.hist(scores, bins=min(10, len(set(scores))), alpha=0.7, color='skyblue', edgecolor='black')
            ax.set_xlabel("Điểm số")
            ax.set_ylabel("Số lượng bài nộp")
            ax.grid(True, linestyle='--', alpha=0.3)
            plt.tight_layout()
            st.pyplot(fig)
    
    with tab2:
        st.subheader("Chi tiết theo học viên")
        
        # Tạo DataFrame từ dữ liệu
        user_data = []
        for s in submissions:
            # Tìm thông tin học viên
            student_info = next((student for student in students if student["email"] == s["user_email"]), None)
            full_name = student_info["full_name"] if student_info else "Không xác định"
            class_name = student_info["class"] if student_info else "Không xác định"
            
            user_data.append({
                "email": s["user_email"],
                "full_name": full_name,
                "class": class_name,
                "submission_id": s["id"],
                "timestamp": datetime.fromtimestamp(s["timestamp"]).strftime("%H:%M:%S %d/%m/%Y"),
                "score": s["score"],
                "max_score": max_possible,
                "percent": f"{(s['score']/max_possible*100):.1f}%"
            })
        
        if user_data:
            df_users = pd.DataFrame(user_data)
            
            # Lọc theo email hoặc lớp
            col1, col2 = st.columns(2)
            with col1:
                user_filter = st.selectbox(
                    "Chọn học viên để xem chi tiết:",
                    options=["Tất cả"] + sorted(list(set([u["email"] for u in user_data]))),
                    key="user_filter_tab2"
                )
            
            with col2:
                class_filter = st.selectbox(
                    "Lọc theo lớp:",
                    options=["Tất cả"] + sorted(list(set([u["class"] for u in user_data if u["class"] != "Không xác định"]))),
                    key="class_filter_tab2"
                )
            
            # Áp dụng bộ lọc
            df_filtered = df_users
            
            if user_filter != "Tất cả":
                df_filtered = df_filtered[df_filtered["email"] == user_filter]
            
            if class_filter != "Tất cả":
                df_filtered = df_filtered[df_filtered["class"] == class_filter]
            
            # Hiển thị bảng
            st.dataframe(
                df_filtered.sort_values(by="timestamp", ascending=False),
                use_container_width=True,
                hide_index=True,
                column_order=["full_name", "class", "email", "score", "percent", "timestamp"]
            )
            
            # Xem chi tiết một bài nộp cụ thể
            if user_filter != "Tất cả":
                submission_ids = df_filtered["submission_id"].tolist()
                if submission_ids:
                    selected_submission = st.selectbox(
                        "Chọn bài nộp để xem chi tiết:",
                        options=submission_ids,
                        key="submission_id_select"
                    )
                    
                    # Tìm bài nộp được chọn
                    submission = next((s for s in submissions if s["id"] == selected_submission), None)
                    if submission:
                        st.subheader(f"Chi tiết bài nộp #{selected_submission}")
                        
                        total_correct = 0
                        total_questions = len(questions)
                        student_detail_data = []
                        
                        # Hiển thị câu trả lời chi tiết
                        for q in questions:
                            q_id = str(q["id"])
                            st.write(f"**Câu {q['id']}: {q['question']}**")
                            
                            # Đáp án người dùng
                            user_ans = submission["responses"].get(q_id, [])
                            expected = [q["answers"][i - 1] for i in q["correct"]]
                            
                            # Kiểm tra đúng/sai
                            is_correct = set(user_ans) == set(expected)
                            if is_correct:
                                total_correct += 1
                            
                            # Thu thập dữ liệu chi tiết
                            student_detail_data.append({
                                "Câu hỏi": f"Câu {q['id']}: {q['question']}",
                                "Đáp án của học viên": ", ".join(user_ans) if user_ans else "Không trả lời",
                                "Đáp án đúng": ", ".join(expected),
                                "Kết quả": "Đúng" if is_correct else "Sai",
                                "Điểm": q["score"] if is_correct else 0
                            })
                            
                            # Hiển thị đáp án của người dùng
                            st.write("Đáp án của học viên:")
                            if not user_ans:
                                st.write("- Không trả lời")
                            else:
                                for ans in user_ans:
                                    st.write(f"- {ans}")
                            
                            # Hiển thị kết quả
                            if is_correct:
                                st.success(f"✅ Đúng (+{q['score']} điểm)")
                            else:
                                st.error("❌ Sai (0 điểm)")
                                st.write("Đáp án đúng:")
                                for ans in expected:
                                    st.write(f"- {ans}")
                            
                            st.divider()
                            
                        # Hiển thị thống kê tổng hợp
                        st.subheader("Tổng kết")
                        col1, col2, col3 = st.columns(3)
                        col1.metric("Số câu đúng", f"{total_correct}/{total_questions}")
                        col2.metric("Điểm số", f"{submission['score']}/{max_possible}")
                        col3.metric("Tỷ lệ đúng", f"{(total_correct/total_questions*100):.1f}%")
                        
                        # Tạo DataFrame chi tiết
                        df_student_detail = pd.DataFrame(student_detail_data)
                        
                        # Xuất báo cáo chi tiết
                        st.write("### Xuất báo cáo chi tiết")
                        
                        # Người dùng và thông tin
                        student_info = next((student for student in students if student["email"] == submission["user_email"]), None)
                        student_name = student_info["full_name"] if student_info else "Không xác định"
                        student_class = student_info["class"] if student_info else "Không xác định"
                        
                        # Tạo báo cáo dạng DOCX
                        try:
                            docx_buffer = create_student_report_docx(
                                student_name,
                                submission["user_email"],
                                student_class,
                                submission,
                                questions,
                                max_possible
                            )
                            
                            st.markdown(
                                get_download_link_docx(docx_buffer, 
                                                     f"bao_cao_{student_name}_{submission['id']}.docx", 
                                                     "Tải xuống báo cáo chi tiết (DOCX)"), 
                                unsafe_allow_html=True
                            )
                        except Exception as e:
                            st.error(f"Không thể tạo báo cáo: {str(e)}")
    
    with tab3:
        st.subheader("Phân tích theo câu hỏi")
        
        # Thống kê tỷ lệ đúng/sai cho từng câu hỏi
        question_stats = {}
        
        for q in questions:
            q_id = str(q["id"])
            correct_count = 0
            wrong_count = 0
            skip_count = 0
            
            for s in submissions:
                user_ans = s["responses"].get(q_id, [])
                expected = [q["answers"][i - 1] for i in q["correct"]]
                
                if not user_ans:
                    skip_count += 1
                elif set(user_ans) == set(expected):
                    correct_count += 1
                else:
                    wrong_count += 1
            
            question_stats[q_id] = {
                "question": q["question"],
                "correct": correct_count,
                "wrong": wrong_count,
                "skip": skip_count,
                "total": correct_count + wrong_count + skip_count,
                "correct_rate": correct_count / (correct_count + wrong_count + skip_count) if (correct_count + wrong_count + skip_count) > 0 else 0
            }
        
        # DataFrame thống kê câu hỏi
        df_questions = pd.DataFrame([
            {
                "Câu hỏi ID": q_id,
                "Nội dung": stats["question"],
                "Số lượng đúng": stats["correct"],
                "Số lượng sai": stats["wrong"],
                "Bỏ qua": stats["skip"],
                "Tổng số làm": stats["total"],
                "Tỷ lệ đúng (%)": f"{stats['correct_rate']*100:.1f}%"
            }
            for q_id, stats in question_stats.items()
        ])
        
        # Vẽ biểu đồ tỷ lệ đúng theo từng câu hỏi
        q_ids = list(question_stats.keys())
        correct_rates = [question_stats[q_id]["correct_rate"] * 100 for q_id in q_ids]
        
        # Tạo biểu đồ với kích thước nhỏ hơn
        fig, ax = plt.subplots(figsize=(10, 4))
        bars = ax.bar(q_ids, correct_rates, color='skyblue')
        
        # Thêm nhãn giá trị
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 1,
                    f'{height:.1f}%', ha='center', va='bottom', fontsize=9)
        
        ax.set_ylim(0, 105)  # Giới hạn trục y từ 0-100%
        ax.set_xlabel("Câu hỏi")
        ax.set_ylabel("Tỷ lệ đúng (%)")
        ax.set_title("Tỷ lệ trả lời đúng theo từng câu hỏi")
        ax.grid(axis='y', linestyle='--', alpha=0.7)
        plt.tight_layout()
        st.pyplot(fig)
        
        # Hiển thị bảng thống kê
        st.dataframe(df_questions, use_container_width=True, hide_index=True)
        
        # Chi tiết từng câu hỏi
        selected_question = st.selectbox(
            "Chọn câu hỏi để xem chi tiết:",
            options=[(f"Câu {q_id}: {question_stats[q_id]['question']}") for q_id in q_ids],
            key="question_select_tab3"
        )
        
        if selected_question:
            q_id = selected_question.split(":")[0].replace("Câu ", "").strip()
            q_data = question_stats[q_id]
            q_detail = next((q for q in questions if str(q["id"]) == q_id), None)
            
            if q_detail:
                st.write(f"**{selected_question}**")
                
                # Hiển thị thống kê
                col1, col2, col3, col4 = st.columns(4)
                col1.metric("✅ Đúng", q_data["correct"])
                col2.metric("❌ Sai", q_data["wrong"])
                col3.metric("⏭️ Bỏ qua", q_data["skip"])
                col4.metric("📊 Tỷ lệ đúng", f"{q_data['correct_rate']*100:.1f}%")
                
                # Tạo biểu đồ tròn nhỏ hơn
                fig, ax = plt.subplots(figsize=(6, 4))
                
                # Sử dụng biểu đồ đơn giản với nhãn và tỷ lệ bên ngoài
                labels = ['Đúng', 'Sai', 'Bỏ qua']
                sizes = [q_data["correct"], q_data["wrong"], q_data["skip"]]
                colors = ['#4CAF50', '#F44336', '#9E9E9E']
                
                # Chỉ hiển thị phần trăm nếu giá trị > 0
                patches, texts, autotexts = ax.pie(
                    sizes, 
                    labels=None,  # Không hiển thị nhãn trên biểu đồ
                    colors=colors, 
                    autopct=lambda p: f'{p:.1f}%' if p > 0 else '',
                    startangle=90,
                    pctdistance=0.85  # Đặt phần trăm gần hơn với trung tâm
                )
                
                # Thiết lập kích thước font nhỏ hơn
                for autotext in autotexts:
                    autotext.set_fontsize(9)
                
                # Thêm chú thích bên ngoài biểu đồ
                ax.legend(labels, loc="upper right", fontsize=9)
                
                # Vẽ vòng tròn trắng ở giữa
                centre_circle = plt.Circle((0, 0), 0.5, fc='white')
                ax.add_patch(centre_circle)
                
                ax.axis('equal')  # Giữ tỷ lệ vòng tròn
                plt.tight_layout()
                st.pyplot(fig)
                
                # Hiển thị đáp án đúng
                st.write("**Đáp án đúng:**")
                for i in q_detail["correct"]:
                    st.write(f"- {q_detail['answers'][i-1]}")
    
    with tab4:
        st.subheader("Danh sách học viên")
        
        if not students:
            st.info("Chưa có học viên nào đăng ký")
        else:
            # Chuẩn bị dữ liệu
            student_data = []
            for student in students:
                # Tìm tất cả bài nộp của học viên
                student_submissions = [s for s in submissions if s["user_email"] == student["email"]]
                submission_count = len(student_submissions)
                
                # Tìm điểm cao nhất
                max_student_score = max([s["score"] for s in student_submissions]) if student_submissions else 0
                
                # Thời gian đăng ký
                registration_date = format_date(student.get("registration_date"))
                
                student_data.append({
                    "full_name": student["full_name"],
                    "email": student["email"],
                    "class": student["class"],
                    "registration_date": registration_date,
                    "submission_count": submission_count,
                    "max_score": max_student_score,
                    "max_possible": max_possible,
                    "percent": f"{(max_student_score/max_possible*100):.1f}%" if max_possible > 0 else "N/A"
                })
            
            # DataFrame cho danh sách học viên
            df_students_list = pd.DataFrame([
                {
                    "Họ và tên": s["full_name"],
                    "Email": s["email"],
                    "Lớp": s["class"],
                    "Ngày đăng ký": s["registration_date"],
                    "Số lần làm bài": s["submission_count"],
                    "Điểm cao nhất": s["max_score"],
                    "Điểm tối đa": s["max_possible"],
                    "Tỷ lệ đúng": s["percent"]
                } for s in student_data
            ])
            
            # Lọc theo lớp
            class_filter = st.selectbox(
                "Lọc theo lớp:",
                options=["Tất cả"] + sorted(list(set([s["class"] for s in student_data if s["class"]]))),
                key="class_filter_tab4"
            )
            
            df_students = pd.DataFrame(student_data)
            
            if class_filter != "Tất cả":
                df_students = df_students[df_students["class"] == class_filter]
            
            # Sắp xếp theo tên
            df_students = df_students.sort_values(by="full_name")
            
            # Hiển thị bảng
            st.dataframe(
                df_students,
                use_container_width=True,
                hide_index=True,
                column_order=["full_name", "class", "email", "submission_count", "max_score", "percent", "registration_date"]
            )
            
            # Thống kê theo lớp
            st.subheader("Thống kê theo lớp")
            
            # Nhóm theo lớp
            class_stats = df_students.groupby("class").agg({
                "email": "count",
                "submission_count": "sum",
                "max_score": "mean"
            }).reset_index()
            
            class_stats.columns = ["Lớp", "Số học viên", "Tổng số bài nộp", "Điểm trung bình"]
            class_stats["Điểm trung bình"] = class_stats["Điểm trung bình"].round(2)
            
            # DataFrame thống kê lớp
            df_class_stats = class_stats.copy()
            
            st.dataframe(
                class_stats,
                use_container_width=True,
                hide_index=True
            )
            
            # Biểu đồ cột nhỏ hơn cho số học viên theo lớp
            fig, ax = plt.subplots(figsize=(10, 4))
            ax.bar(class_stats["Lớp"], class_stats["Số học viên"], color='skyblue')
            ax.set_xlabel("Lớp")
            ax.set_ylabel("Số học viên")
            ax.set_title("Số học viên theo lớp")
            plt.xticks(rotation=45, ha='right')
            plt.tight_layout()
            st.pyplot(fig)
    
    with tab5:
        st.subheader("Xuất báo cáo")
        
        # Hiển thị các loại báo cáo có thể xuất
        st.write("### 1. Báo cáo tất cả bài nộp")
        
        try:
            # DOCX
            docx_buffer = dataframe_to_docx(df_all_submissions, "Báo cáo tất cả bài nộp", "bao_cao_tat_ca_bai_nop.docx")
            st.markdown(get_download_link_docx(docx_buffer, "bao_cao_tat_ca_bai_nop.docx", 
                                        "Tải xuống báo cáo (DOCX)"), unsafe_allow_html=True)
        except Exception as e:
            st.error(f"Lỗi khi tạo DOCX: {str(e)}")
        
        st.write("### 2. Báo cáo thống kê câu hỏi")
        
        try:
            # DOCX
            docx_buffer = dataframe_to_docx(df_questions, "Báo cáo thống kê câu hỏi", "bao_cao_thong_ke_cau_hoi.docx")
            st.markdown(get_download_link_docx(docx_buffer, "bao_cao_thong_ke_cau_hoi.docx", 
                                        "Tải xuống báo cáo (DOCX)"), unsafe_allow_html=True)
        except Exception as e:
            st.error(f"Lỗi khi tạo DOCX: {str(e)}")
        
        st.write("### 3. Báo cáo danh sách học viên")
        
        try:
            # DOCX
            docx_buffer = dataframe_to_docx(df_students_list, "Báo cáo danh sách học viên", "bao_cao_danh_sach_hoc_vien.docx")
            st.markdown(get_download_link_docx(docx_buffer, "bao_cao_danh_sach_hoc_vien.docx", 
                                        "Tải xuống báo cáo (DOCX)"), unsafe_allow_html=True)
        except Exception as e:
            st.error(f"Lỗi khi tạo DOCX: {str(e)}")
        
        st.write("### 4. Báo cáo thống kê theo lớp")
        
        try:
            # DOCX
            docx_buffer = dataframe_to_docx(df_class_stats, "Báo cáo thống kê theo lớp", "bao_cao_thong_ke_lop.docx")
            st.markdown(get_download_link_docx(docx_buffer, "bao_cao_thong_ke_lop.docx", 
                                        "Tải xuống báo cáo (DOCX)"), unsafe_allow_html=True)
        except Exception as e:
            st.error(f"Lỗi khi tạo DOCX: {str(e)}")
        
        st.write("### 5. Báo cáo tổng hợp (Excel)")
        
        try:
            # Chuẩn bị danh sách DataFrame và tên sheet
            dfs = [df_all_submissions, df_questions, df_students_list, df_class_stats]
            sheet_names = ["Tất cả bài nộp", "Thống kê câu hỏi", "Danh sách học viên", "Thống kê lớp"]
            
            # Hiển thị link tải xuống
            st.markdown(export_to_excel(dfs, sheet_names, "bao_cao_tong_hop.xlsx"), unsafe_allow_html=True)
            
        except Exception as e:
            st.error(f"Lỗi khi tạo file Excel: {str(e)}")